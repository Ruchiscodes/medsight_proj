 

import os
import io
import re
import json
import streamlit as st
import numpy as np
import pydicom
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
import requests
import google.generativeai as genai

# ---------------- PubMed Tool ----------------
def search_pubmed(query: str) -> str:
    """Searches PubMed for top 3 results and returns as JSON string."""
    st.write(f"🔬 Searching PubMed for: `{query}`...")
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    search_url = f"{base_url}esearch.fcgi?db=pubmed&term={query}&retmode=json&retmax=3"
    try:
        resp = requests.get(search_url, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        ids = data.get("esearchresult", {}).get("idlist", [])
        if not ids:
            return "No relevant articles found on PubMed."
        ids_str = ",".join(ids)
        summary_url = f"{base_url}esummary.fcgi?db=pubmed&id={ids_str}&retmode=json"
        summary_resp = requests.get(summary_url, timeout=10)
        summary_resp.raise_for_status()
        summary_data = summary_resp.json()
        results = []
        for uid in ids:
            article = summary_data['result'][uid]
            results.append({
                "title": article.get('title', 'N/A'),
                "authors": [a.get('name', 'N/A') for a in article.get('authors', [])],
                "journal": article.get('fulljournalname', 'N/A'),
                "pub_date": article.get('pubdate', 'N/A'),
                "pmid": uid,
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{uid}/"
            })
        return json.dumps(results, indent=2)
    except Exception as e:
        return f"PubMed search failed: {e}"

        
         

# ---------------- DICOM Helpers ----------------
def get_all_dicom_metadata(dicom_file: pydicom.FileDataset) -> str:
    metadata = "--- Full DICOM Header Metadata ---\n"
    for tag in dicom_file.iterall():
        if tag.keyword != "PixelData":
            try:
                metadata += f"{tag.name} ({tag.tag}): {tag.value}\n"
            except Exception:
                metadata += f"{tag.name} ({tag.tag}): [Unreadable]\n"
    metadata += "---------------------------------\n"
    return metadata

from pydicom.pixel_data_handlers.util import apply_modality_lut

def handle_dicom_file(uploaded_file):
    try:
        dicom_bytes = io.BytesIO(uploaded_file.getvalue())
        dicom_data = pydicom.dcmread(dicom_bytes)
        full_metadata_text = get_all_dicom_metadata(dicom_data)
        images = []

        # Decompress if needed using pylibjpeg / apply LUT
        try:
            pixel_array_full = apply_modality_lut(dicom_data.pixel_array, dicom_data)
        except Exception as e:
            st.error(f"Error decompressing pixel data: {e}")
            return [], None

        # Handle multi-frame DICOMs
        if hasattr(dicom_data, "NumberOfFrames") and dicom_data.NumberOfFrames > 1:
            st.info(f"Multi-frame DICOM detected ({dicom_data.NumberOfFrames} frames).")
            frame_indices = [0, dicom_data.NumberOfFrames // 2, dicom_data.NumberOfFrames - 1]
            pixel_arrays = pixel_array_full  # already decompressed
        else:
            frame_indices = [0]
            pixel_arrays = [pixel_array_full]

        for i in frame_indices:
            pixel_array = pixel_arrays[i]
            pixel_array = pixel_array.astype(float)
            rescaled_array = (np.maximum(pixel_array, 0) / pixel_array.max()) * 255
            final_array = np.uint8(rescaled_array)
            image = PILImage.fromarray(final_array)
            if image.mode != "RGB":
                image = image.convert("RGB")
            images.append(image)

        return images, full_metadata_text

    except Exception as e:
        st.error(f"Error processing DICOM file: {e}")
        return [], None


# ---------------- PDF Helper ----------------
def handle_pdf_file(uploaded_file):
    try:
        pdf_bytes = uploaded_file.getvalue()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text_content = ""
        images = []
        for page in doc:
            text_content += page.get_text()
        for page_num in range(len(doc)):
            imgs = doc.get_page_images(page_num)
            if imgs:
                xref = imgs[0][0]
                base_img = doc.extract_image(xref)
                img_bytes = base_img["image"]
                img = PILImage.open(io.BytesIO(img_bytes))
                if img.mode != "RGB":
                    img = img.convert("RGB")
                images.append(img)
                break
        return images, text_content
    except Exception as e:
        st.error(f"Error processing PDF: {e}")
        return [], None

# ---------------- Streamlit UI ----------------
st.title("🏥 MedSight AI Pro: Advanced Medical Analysis Agent")

# Sidebar: API Key
st.sidebar.title("⚙️ Configuration")
api_key = st.sidebar.text_input("Enter Google AI API Key:", type="password")
if st.sidebar.button("Set API Key") and api_key:
    st.session_state.GOOGLE_API_KEY = api_key
    st.sidebar.success("✅ API Key configured!")

if "GOOGLE_API_KEY" not in st.session_state:
    st.warning("Please configure your Google AI API key in the sidebar to continue.")
    st.stop()

genai.configure(api_key=st.session_state["GOOGLE_API_KEY"])

# File uploader
uploaded_file = st.file_uploader(
    "Upload a Medical File (DICOM, PDF, JPG, PNG)",
    type=["dcm", "dicom", "pdf", "jpg", "jpeg", "png"]
)

medical_prompt = """
You are a world-class radiology expert AI.
Generate a **two-part medical report**:
1️⃣ Professional Report: Include quantitative findings, differential diagnoses, and recommendations with PubMed citations.
2️⃣ Patient Summary: Write in simple terms with emojis and next steps.
"""

if uploaded_file:
    ext = uploaded_file.name.lower().split(".")[-1]
    images, extra_text = [], ""

    if ext in ["dcm", "dicom"]:
        images, extra_text = handle_dicom_file(uploaded_file)
    elif ext == "pdf":
        images, extra_text = handle_pdf_file(uploaded_file)
    else:
        images = [PILImage.open(uploaded_file)]

    if images:
        st.image(images, caption=[f"Image {i+1}" for i in range(len(images))], width='stretch')
    if extra_text:
        with st.expander("Extracted Text for Context"):
            st.text(extra_text)

    if st.button("🔍 Analyze File with Gemini"):
        with st.spinner("Analyzing with Gemini model..."):
            try:
                # Correct: Create GenerativeModel object
                model = genai.GenerativeModel(model_name="gemini-2.5-pro")

                # Prepare prompt and images
                inputs = [medical_prompt + "\n" + (extra_text or "")]
                inputs.extend(images)  # Gemini can handle images in this format

                response = model.generate_content(inputs)
                result_text = response.text
                st.session_state["analysis_result"] = result_text

                # # Optional: PubMed references
                # keywords = "radiology " + (extra_text[:100] if extra_text else "")
                # pubmed_refs = search_pubmed(keywords)
                # st.markdown("### 🔬 PubMed References")
                # st.json(json.loads(pubmed_refs) if pubmed_refs.startswith("[") else pubmed_refs)
                # Optional: PubMed references (short format)
                keywords = "radiology " + (extra_text[:100] if extra_text else "")
                pubmed_refs_json = search_pubmed(keywords)
                pubmed_refs_short = []

                try:
                    pubmed_refs_data = json.loads(pubmed_refs_json)
                    for ref in pubmed_refs_data:
                        title = ref.get("title", "N/A")
                        url = ref.get("url", "")
                        pubmed_refs_short.append(f"- [{title}]({url})")
                except Exception:
                    pubmed_refs_short.append(pubmed_refs_json)

                # Display at the end
                st.markdown("### 🔬 PubMed References")
                st.markdown("\n".join(pubmed_refs_short))

            except Exception as e:
                st.error(f"An error occurred during analysis: {e}")

# ---------------- DOCX Report ----------------
def create_docx(text, imgs):
    doc = Document()
    doc.add_heading("Medical Analysis Report", 0)
    if imgs:
        doc.add_heading("Analyzed Images", level=1)
        for i, img in enumerate(imgs):
            doc.add_paragraph(f"Image {i+1}")
            img_io = io.BytesIO()
            img.save(img_io, format="PNG")
            img_io.seek(0)
            doc.add_picture(img_io, width=Inches(5.0))
    doc.add_paragraph(text)
    return doc

def get_docx_bytes(text, imgs):
    doc = create_docx(text, imgs)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

if st.session_state.get("analysis_result"):
    st.markdown("### 📋 Analysis Report")
    st.markdown(st.session_state["analysis_result"], unsafe_allow_html=True)
    st.download_button(
        "⬇️ Download Report (DOCX)",
        data=get_docx_bytes(st.session_state["analysis_result"], images),
        file_name="Medical_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
